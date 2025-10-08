import matplotlib.pyplot as plt
from collections import Counter
import numpy as np
import pycountry

from collections import Counter
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document

def compare_histograms(items_list1, items_list2, title, filename,
                       label1='List 1', label2='List 2',
                       scale_counts=True, order_by_count=False):
    # Count frequencies
    counter1 = Counter(items_list1)
    counter2 = Counter(items_list2)

    # Get all unique labels
    all_labels = list(set(counter1.keys()).union(set(counter2.keys())))

    # Total counts for normalization
    total1 = sum(counter1.values())
    total2 = sum(counter2.values())

    if scale_counts:
        # Frequencies (proportions)
        freq1_dict = {label: counter1.get(label, 0) / total1 for label in all_labels}
        freq2_dict = {label: counter2.get(label, 0) / total2 for label in all_labels}
    else:
        # Raw counts
        freq1_dict = {label: counter1.get(label, 0) for label in all_labels}
        freq2_dict = {label: counter2.get(label, 0) for label in all_labels}

    # Optionally reorder labels by highest total count/frequency
    if order_by_count:
        all_labels = sorted(
            all_labels,
            key=lambda lbl: freq1_dict[lbl] + freq2_dict[lbl],
            reverse=True
        )
    else:
        all_labels = sorted(all_labels)

    # Get ordered frequencies
    freq1 = [freq1_dict[label] for label in all_labels]
    freq2 = [freq2_dict[label] for label in all_labels]

    # Set positions for bars
    x = np.arange(len(all_labels))
    width = 0.4  # width of the bars

    # Create the grouped bar chart
    plt.figure(figsize=(24, 8))
    plt.bar(x - width/2, freq1, width, label=label1, color='skyblue')
    plt.bar(x + width/2, freq2, width, label=label2, color='salmon')

    # Formatting
    plt.ylabel('Frequency' if scale_counts else 'Count', fontsize=18)
    plt.title(title, fontsize=22)
    plt.xticks(x, all_labels, rotation=45, ha='right', fontsize=18)
    plt.yticks(fontsize=18)
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.legend(fontsize=18)
    plt.tight_layout()

    # Save and show
    plt.savefig(filename, dpi=300, facecolor='white', edgecolor='none')
    plt.show()


import matplotlib.pyplot as plt
import seaborn as sns

def line_time_plot(
    df_articles_with_results,
    nyt_mask,
    variable,
    title,
    filename,
    variable_label=None,
    smooth_window=3  # rolling average window; set to None to disable
):
    """
    Plot two line charts of article counts per year, grouped by a categorical variable:
    one for NYT and one for other publishers, with consistent colors. Optionally smooth lines.

    Parameters
    ----------
    df_articles_with_results : pandas.DataFrame
        DataFrame containing a 'year' column and the specified categorical variable.
    nyt_mask : pandas.Series[bool]
        Boolean mask selecting NYT articles.
    variable : str
        Column name of the categorical variable.
    title : str
        Base title for the plots.
    variable_label : str, optional
        Label to use for the variable in the legend. Defaults to `variable` if None.
    smooth_window : int or None, optional
        Rolling window size for smoothing lines. If None, no smoothing is applied.
    """
    if variable_label is None:
        variable_label = variable

    # Determine all categories
    all_categories = df_articles_with_results[variable].dropna().unique()

    # Use a clean, pretty color palette
    palette = sns.color_palette("tab10", n_colors=len(all_categories))
    colors = {cat: palette[i] for i, cat in enumerate(all_categories)}

    fig, axes = plt.subplots(2, 1, figsize=(12, 10), sharex=True)

    for ax, mask, subtitle in zip(
        axes,
        [nyt_mask, ~nyt_mask],
        [f"NYT: {title}", f"Other Publishers: {title}"]
    ):
        grouped = (
            df_articles_with_results[mask]
            .groupby(["year", variable])
            .size()
            .unstack(fill_value=0)
        )

        # Reindex columns to match global category order
        grouped = grouped.reindex(columns=all_categories, fill_value=0)

        # Optionally smooth with rolling average
        if smooth_window is not None and smooth_window > 1:
            grouped = grouped.rolling(window=smooth_window, min_periods=1).mean()

        # Plot each category as a line
        for cat in grouped.columns:
            ax.plot(
                grouped.index,
                grouped[cat],
                label=cat,
                color=colors[cat],
                linewidth=2,
                alpha=0.9
            )

        ax.set_ylabel("Frequency")
        ax.set_title(subtitle)
        ax.grid(True, linestyle="--", alpha=0.7)
        ax.legend(title=variable_label, bbox_to_anchor=(1.05, 1), loc="upper left", ncol=2)

    axes[-1].set_xlabel("Year")
    plt.tight_layout()

    # Save and show
    plt.savefig(filename, dpi=300, facecolor='white', edgecolor='none')
    plt.show()



# Function to split states and countries
def split_locations(locations):
    us_states = {
        'Alabama', 'Alaska', 'Arizona', 'Arkansas', 'California', 'Colorado', 'Connecticut',
        'Delaware', 'Florida', 'Georgia', 'Hawaii', 'Idaho', 'Illinois', 'Indiana', 'Iowa',
        'Kansas', 'Kentucky', 'Louisiana', 'Maine', 'Maryland', 'Massachusetts', 'Michigan',
        'Minnesota', 'Mississippi', 'Missouri', 'Montana', 'Nebraska', 'Nevada', 'New Hampshire',
        'New Jersey', 'New Mexico', 'New York', 'North Carolina', 'North Dakota', 'Ohio',
        'Oklahoma', 'Oregon', 'Pennsylvania', 'Rhode Island', 'South Carolina', 'South Dakota',
        'Tennessee', 'Texas', 'Utah', 'Vermont', 'Virginia', 'Washington', 'West Virginia',
        'Wisconsin', 'Wyoming'
    }

    world_countries = [c.name for c in pycountry.countries if c.name not in ["Jersey", "Georgia"]] + ["UK", "Britain", "England", "South Korea", "North Korea", "Iran", "Syria", "Venezuela", "Russia", "Taiwan", "Saint Martin", "Iran", "Lao", "Moldova", "Tanzania", "Vatican", "Venezuela", "Korea", "Turkey", "Czechoslovakia", "Congo"]

    states = []
    countries = []
    other = []

    for loc in locations:
        state_match = next((state for state in us_states if state in loc), None)
        country_match = next((country for country in world_countries if country in loc), None)

        if state_match:
            states.append(state_match)
        elif country_match:
            countries.append(country_match)
        else:
            other.append(loc)

    return states, countries, other


# Flatten and filter location lists
#def flatten_and_filter(locations_column):
 #   return [item for sublist in locations_column for item in sublist if item != "NA"]

from itertools import chain

# Flatten and filter location lists
#def flatten_and_filter(series, exclude="NA"):
#    """Flatten lists in a pandas Series and filter out unwanted values."""
#    return [item for item in chain.from_iterable(series.dropna()) if item != exclude]

def flatten(series):
    """Flatten lists in a pandas Series"""
    return [item for item in chain.from_iterable(series.dropna())]


import re

def extract_region_mentions(text, regions=None):
    """
    Extract mentions of specified U.S. regions from the given text.

    Parameters
    ----------
    text : str
        The input text to search.
    regions : list of str, optional
        A list of region names to search for. If None, defaults to:
        ['Northeast', 'South', 'Midwest', 'Southwest', 'West']

    Returns
    -------
    list of str
        Mentions of regions found in the text, preserving order of appearance.
        Special rule: mentions of 'Mid-Atlantic' are mapped to 'Northeast'.
    """

    mentions = []
    if regions is None:
        regions = ["Northeast", "South", "Midwest", "Southwest", "West"]

        # Special rule: Mid-Atlantic counts as Northeast
        if re.search(r"\bMid-?Atlantic\b", text): # flags=re.IGNORECASE
            mentions.append("Northeast")

        if re.search(r"\bSouthern states\b", text): # flags=re.IGNORECASE
            mentions.append("South")

    for region in regions:
        pattern = r'\b' + re.escape(region) + r'\b'
        if re.search(pattern, text): # flags=re.IGNORECASE
            mentions.append(region)

    return mentions


import json
import time
import pandas as pd
from openai import OpenAI

client = OpenAI()


# --- Create batch ---
def create_batch(df, system_prompt, response_format, model, temperature, reasoning_effort, batch_input_filename="batch_input.jsonl"):
    """
    Convert DataFrame into a JSONL file suitable for the OpenAI batch API.
    Each row becomes one request.
    """
    with open(batch_input_filename, "w") as f:
        for _, row in df.iterrows():
            user_prompt = f"Title: {row['title']}\n\nBody: {row['body']}"
            request = {
                "custom_id": f"req-{row.name}",  # track back to row
                "method": "POST",
                "url": "/v1/chat/completions",
                "body": {
                    "model": model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "response_format": response_format,
                    "temperature": temperature,
                    "reasoning_effort": reasoning_effort
                },
            }
            f.write(json.dumps(request) + "\n")
    return batch_input_filename


def submit_batch(batch_input_filename):
    """
    Upload batch file and create a batch job.
    """
    # Step 1: Upload file
    with open(batch_input_filename, "rb") as f:
        uploaded_file = client.files.create(file=f, purpose="batch")

    # Step 2: Create batch using uploaded file id
    batch = client.batches.create(
        input_file_id=uploaded_file.id,
        endpoint="/v1/chat/completions",
        completion_window="24h"
    )

    print(f"Batch submitted. ID: {batch.id}")
    return batch.id


def retrieve_batch_results(batch_id, output_filename="batch_output.jsonl"):
    """
    Download completed batch results.
    """
    batch = client.batches.retrieve(batch_id)
    if batch.status != "completed":
        print(f"Batch not ready yet. Status: {batch.status}")
        return None

    output_file = client.files.retrieve(batch.output_file_id)
    result_content = client.files.content(output_file.id).text

    with open(output_filename, "w") as f:
        f.write(result_content)

    print(f"Results saved to {output_filename}")
    return output_filename


pricing_dict = {
    'gpt-5-mini': {'Input': 0.25, 'Cached input': 0.025, 'Output': 2.00},
    'gpt-5': {'Input': 1.25, 'Cached input': 0.125, 'Output': 10.00},
    'gpt-4.1':  {'Input': 2.00, 'Cached input': 0.5, 'Output': 8.00}
}

def parse_batch_results(output_filename, model, pricing_dict=pricing_dict):
    """
    Parse JSONL batch output file into DataFrame of responses,
    including token usage and cached inputs for cost calculation.
    Uses pricing_dict to compute estimated cost (values are $ per 1M tokens).
    """
    responses = []
    total_prompt_tokens = 0
    total_completion_tokens = 0
    total_cached_tokens = 0

    with open(output_filename, "r") as f:
        for line in f:
            item = json.loads(line)
            custom_id = item["custom_id"]

            body = item["response"]["body"]
            choice = json.loads(body["choices"][0]["message"]["content"])

            # usage info
            usage = body.get("usage", {})
            prompt_tokens = usage.get("prompt_tokens", 0)
            completion_tokens = usage.get("completion_tokens", 0)
            cached_tokens = (
                usage.get("prompt_tokens_details", {})
                .get("cached_tokens", 0)
            )

            total_prompt_tokens += prompt_tokens
            total_completion_tokens += completion_tokens
            total_cached_tokens += cached_tokens

             # Flatten choice dict into the response row
            row = {
                "custom_id": custom_id,
                "prompt_tokens": prompt_tokens,
                "completion_tokens": completion_tokens,
                "cached_tokens": cached_tokens,
                "total_tokens": usage.get("total_tokens", prompt_tokens + completion_tokens),
            }
            # Add each key from choice dict as a separate column
            row.update(choice)
            responses.append(row)

    df = pd.DataFrame(responses)

    # === Token summary ===
    print("=== Token Usage Summary ===")
    print(f"Prompt tokens:     {total_prompt_tokens}")
    print(f"Completion tokens: {total_completion_tokens}")
    print(f"Cached tokens:     {total_cached_tokens}")
    print(f"Total tokens:      {total_prompt_tokens + total_completion_tokens}")

    # === Cost calculation ===
    if model not in pricing_dict:
        raise ValueError(f"Model {model} not found in pricing_dict")

    # Convert $ per 1M tokens to $ per token
    rates = pricing_dict[model]
    prompt_rate = rates["Input"] / 1_000_000
    cached_rate = rates["Cached input"] / 1_000_000
    completion_rate = rates["Output"] / 1_000_000

    cost = (
        (total_prompt_tokens - total_cached_tokens) * prompt_rate +
        (total_cached_tokens * cached_rate) +
        (total_completion_tokens * completion_rate)
    ) / 2 # 50% for batch processing

    print(f"Estimated cost for {model}: ${cost:.4f}")

    return df









# archive
import matplotlib.pyplot as plt

def stacked_time_plot(df_articles_with_results, nyt_mask, variable, title):
    """
    Plot two stacked bar charts of article counts per year, grouped by a categorical variable:
    one for NYT and one for other publishers, with consistent color coding.
    """
    # Ensure consistent categories and color mapping
    all_categories = (
        df_articles_with_results[variable]
        .value_counts()
        .index
    )

    color_map = plt.colormaps.get_cmap("Set2").resampled(len(all_categories))
    colors = {cat: color_map(i) for i, cat in enumerate(all_categories)}

    fig, axes = plt.subplots(2, 1, figsize=(12, 10), sharex=True)

    for ax, mask, subtitle in zip(
        axes,
        [nyt_mask, ~nyt_mask],
        [f"NYT: {title}", f"Other Publishers: {title}"],
    ):
        grouped = (
            df_articles_with_results[mask]
            .groupby(["year", variable])
            .size()
            .unstack(fill_value=0)
        )

        # Reorder columns to match global order
        grouped = grouped.reindex(columns=all_categories, fill_value=0)

        grouped.plot(
            kind="bar",
            stacked=True,
            ax=ax,
            color=[colors[cat] for cat in grouped.columns],
        )

        ax.set_ylabel("Frequency")
        ax.set_title(subtitle)
        ax.tick_params(axis="x", rotation=45)
        ax.grid(axis="y", linestyle="--", alpha=0.7)
        ax.legend(title=variable, bbox_to_anchor=(1.05, 1), loc="upper left")

    axes[-1].set_xlabel("Year")
    plt.tight_layout()
    plt.show()

def save_examples_to_folder(var, df_articles_with_results, examples_dir, var_original = None):

    var_abbr = var.removesuffix("_stand_flat")

    base_folder = os.path.join(examples_dir, var_abbr)
    if not os.path.exists(base_folder):
        os.makedirs(base_folder)

    # 2️⃣ Iterate over each unique value in race_ethnicity_stand_flat
    for val in set(df_articles_with_results[var]):
        # Create subfolder for this race/ethnicity if it doesn't exist
        subfolder_path = os.path.join(base_folder, str(val))
        if not os.path.exists(subfolder_path):
            os.makedirs(subfolder_path)

        # 3️⃣ Select 3 New York Times and 3 Other publisher articles for this category
        nyt_articles = df_articles_with_results[
            (df_articles_with_results[var] == val) &
            (df_articles_with_results['publisher'] == "New York Times")
        ].head(3)

        other_articles = df_articles_with_results[
            (df_articles_with_results[var] == val) &
            (df_articles_with_results['publisher'] != "New York Times")
        ].head(3)

        # Combine both selections
        selected_articles = pd.concat([nyt_articles, other_articles], ignore_index=True)

        # 4️⃣ Save each article as a .docx file
        for n, row in enumerate(selected_articles.itertuples(), start=1):
            doc = Document()

            # Add fields to document
            doc.add_heading(str(row.title), level=1)
            doc.add_paragraph(f"Publisher: {row.publisher}")
            doc.add_paragraph(f"Date: {row.date}")
            doc.add_paragraph(f"Section: {row.section}")
            doc.add_paragraph(f"Source File: {row.source_file}")
            if var_original is None:
                doc.add_paragraph(f"Raw {var_abbr} result: {getattr(row, var_abbr)}")
            else:
                doc.add_paragraph(f"Raw {var_original} result: {getattr(row, var_original)}")
            doc.add_paragraph(f"{var} (Standardized): {getattr(row, var)}")
            doc.add_paragraph("\nBody:\n")
            doc.add_paragraph(str(row.body))

            # Create filename
            safe_publisher = row.publisher.replace(" ", "_")
            filename = f"{n}.{var_abbr}_{val}_{safe_publisher}.docx"
            filepath = os.path.join(subfolder_path, filename)

            # Save document
            doc.save(filepath)

    print("✅ Articles saved successfully.")
