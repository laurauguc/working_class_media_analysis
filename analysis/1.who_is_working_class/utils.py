import matplotlib.pyplot as plt
from collections import Counter
import numpy as np
import pycountry
import os
from docx import Document
import re
import seaborn as sns
from itertools import chain
import json
import time
import pandas as pd
from openai import OpenAI

def compare_histograms(items_list1, items_list2, title, filename,
                       label1='List 1', label2='List 2',
                       scale_counts=True, order_by_count=False, ylabel = None):
    # Count frequencies
    counter1 = Counter(items_list1)
    counter2 = Counter(items_list2)

    # Get all unique labels
    all_labels = list(set(counter1.keys()).union(set(counter2.keys())))

    # Total counts for normalization
    total1 = sum(counter1.values())
    total2 = sum(counter2.values())

    if scale_counts:
        # Frequencies as percentages
        freq1_dict = {label: (counter1.get(label, 0) / total1) * 100 for label in all_labels}
        freq2_dict = {label: (counter2.get(label, 0) / total2) * 100 for label in all_labels}
    else:
        # Raw counts
        freq1_dict = {label: counter1.get(label, 0) for label in all_labels}
        freq2_dict = {label: counter2.get(label, 0) for label in all_labels}

    # Optionally reorder labels by highest total frequency or count
    if order_by_count:
        all_labels = sorted(
            all_labels,
            key=lambda lbl: freq1_dict[lbl] + freq2_dict[lbl],
            reverse=True
        )
    else:
        all_labels = sorted(all_labels)

    # Get ordered frequencies
    freq1 = [freq1_dict[label] for label in all_labels]
    freq2 = [freq2_dict[label] for label in all_labels]

    # Set positions for bars
    x = np.arange(len(all_labels))
    width = 0.4  # width of the bars

    # Create the grouped bar chart
    plt.figure(figsize=(24, 8))
    plt.bar(x - width/2, freq1, width, label=label1, color='skyblue')
    plt.bar(x + width/2, freq2, width, label=label2, color='salmon')

    # Formatting
    if ylabel is None:
        ylabel = 'Percentage of Articles (%)' if scale_counts else 'Count'
    plt.ylabel(ylabel, fontsize=18)
    plt.title(title, fontsize=22)
    plt.xticks(x, all_labels, rotation=45, ha='right', fontsize=18)
    plt.yticks(fontsize=18)
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.legend(fontsize=18)
    plt.tight_layout()

    # Save and show
    plt.savefig(filename, dpi=300, facecolor='white', edgecolor='none')
    plt.show()

def line_time_plot(
    df_articles_with_results,
    nyt_mask,
    variable,
    title,
    filename,
    variable_label=None,
    smooth_window=3  # rolling average window; set to None to disable
):
    """
    Plot two line charts showing the percentage of articles per year, grouped by a categorical variable:
    one for NYT and one for other publishers, with consistent colors. Optionally smooth lines.

    Parameters
    ----------
    df_articles_with_results : pandas.DataFrame
        DataFrame containing a 'year' column and the specified categorical variable.
    nyt_mask : pandas.Series[bool]
        Boolean mask selecting NYT articles.
    variable : str
        Column name of the categorical variable.
    title : str
        Base title for the plots.
    variable_label : str, optional
        Label to use for the variable in the legend. Defaults to `variable` if None.
    smooth_window : int or None, optional
        Rolling window size for smoothing lines. If None, no smoothing is applied.
    """
    if variable_label is None:
        variable_label = variable

    # Determine all categories
    all_categories = df_articles_with_results[variable].dropna().unique()

    # Use a clean, pretty color palette
    palette = sns.color_palette("tab10", n_colors=len(all_categories))
    colors = {cat: palette[i] for i, cat in enumerate(all_categories)}

    fig, axes = plt.subplots(2, 1, figsize=(12, 10), sharex=True)

    for ax, mask, subtitle in zip(
        axes,
        [nyt_mask, ~nyt_mask],
        [f"NYT: {title}", f"Other Publishers: {title}"]
    ):
        # Count articles per (year, variable)
        grouped = (
            df_articles_with_results[mask]
            .groupby(["year", variable])
            .size()
            .unstack(fill_value=0)
        )

        # Reindex columns to ensure consistent category order
        grouped = grouped.reindex(columns=all_categories, fill_value=0)

        # Convert to yearly percentages
        grouped = grouped.div(grouped.sum(axis=1), axis=0) * 100

        # Optionally smooth with rolling average
        if smooth_window is not None and smooth_window > 1:
            grouped = grouped.rolling(window=smooth_window, min_periods=1).mean()

        # Plot each category as a line
        for cat in grouped.columns:
            ax.plot(
                grouped.index,
                grouped[cat],
                label=cat,
                color=colors[cat],
                linewidth=2,
                alpha=0.9
            )

        ax.set_ylabel("Percentage of Articles (%)")
        ax.set_title(subtitle)
        ax.grid(True, linestyle="--", alpha=0.7)
        ax.legend(title=variable_label, bbox_to_anchor=(1.05, 1), loc="upper left", ncol=2)

    axes[-1].set_xlabel("Year")
    plt.tight_layout()

    # Save and show
    plt.savefig(filename, dpi=300, facecolor='white', edgecolor='none')
    plt.show()

# Function to split states and countries
def split_locations(locations):
    us_states = {
        'Alabama', 'Alaska', 'Arizona', 'Arkansas', 'California', 'Colorado', 'Connecticut',
        'Delaware', 'Florida', 'Georgia', 'Hawaii', 'Idaho', 'Illinois', 'Indiana', 'Iowa',
        'Kansas', 'Kentucky', 'Louisiana', 'Maine', 'Maryland', 'Massachusetts', 'Michigan',
        'Minnesota', 'Mississippi', 'Missouri', 'Montana', 'Nebraska', 'Nevada', 'New Hampshire',
        'New Jersey', 'New Mexico', 'New York', 'North Carolina', 'North Dakota', 'Ohio',
        'Oklahoma', 'Oregon', 'Pennsylvania', 'Rhode Island', 'South Carolina', 'South Dakota',
        'Tennessee', 'Texas', 'Utah', 'Vermont', 'Virginia', 'Washington', 'West Virginia',
        'Wisconsin', 'Wyoming'
    }

    world_countries = [c.name for c in pycountry.countries if c.name not in ["Jersey", "Georgia"]] + ["UK", "Britain", "England", "South Korea", "North Korea", "Iran", "Syria", "Venezuela", "Russia", "Taiwan", "Saint Martin", "Iran", "Lao", "Moldova", "Tanzania", "Vatican", "Venezuela", "Korea", "Turkey", "Czechoslovakia", "Congo"]

    states = []
    countries = []
    other = []

    for loc in locations:
        state_match = next((state for state in us_states if state in loc), None)
        country_match = next((country for country in world_countries if country in loc), None)

        if state_match:
            states.append(state_match)
        elif country_match:
            countries.append(country_match)
        else:
            other.append(loc)

    return states, countries, other

def flatten(series):
    """Flatten lists in a pandas Series"""
    return [item for item in chain.from_iterable(series.dropna())]

def extract_region_mentions(text, regions=None):
    """
    Extract mentions of specified U.S. regions from the given text.

    Parameters
    ----------
    text : str
        The input text to search.
    regions : list of str, optional
        A list of region names to search for. If None, defaults to:
        ['Northeast', 'South', 'Midwest', 'Southwest', 'West']

    Returns
    -------
    list of str
        Mentions of regions found in the text, preserving order of appearance.
        Special rule: mentions of 'Mid-Atlantic' are mapped to 'Northeast'.
    """

    mentions = []
    if regions is None:
        regions = ["Northeast", "South", "Midwest", "Southwest", "West"]

        # Special rule: Mid-Atlantic counts as Northeast
        if re.search(r"\bMid-?Atlantic\b", text): # flags=re.IGNORECASE
            mentions.append("Northeast")

        if re.search(r"\bSouthern states\b", text): # flags=re.IGNORECASE
            mentions.append("South")

    for region in regions:
        pattern = r'\b' + re.escape(region) + r'\b'
        if re.search(pattern, text): # flags=re.IGNORECASE
            mentions.append(region)

    return mentions

client = OpenAI()

# --- Create batch ---
def create_batch(df, system_prompt, response_format, model, temperature, reasoning_effort, batch_input_filename="batch_input.jsonl"):
    """
    Convert DataFrame into a JSONL file suitable for the OpenAI batch API.
    Each row becomes one request.
    """
    with open(batch_input_filename, "w") as f:
        for _, row in df.iterrows():
            user_prompt = f"Title: {row['title']}\n\nBody: {row['body']}"
            request = {
                "custom_id": f"req-{row.name}",  # track back to row
                "method": "POST",
                "url": "/v1/chat/completions",
                "body": {
                    "model": model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "response_format": response_format,
                    "temperature": temperature,
                    "reasoning_effort": reasoning_effort
                },
            }
            f.write(json.dumps(request) + "\n")
    return batch_input_filename

def submit_batch(batch_input_filename):
    """
    Upload batch file and create a batch job.
    """
    # Step 1: Upload file
    with open(batch_input_filename, "rb") as f:
        uploaded_file = client.files.create(file=f, purpose="batch")

    # Step 2: Create batch using uploaded file id
    batch = client.batches.create(
        input_file_id=uploaded_file.id,
        endpoint="/v1/chat/completions",
        completion_window="24h"
    )

    print(f"Batch submitted. ID: {batch.id}")
    return batch.id

def retrieve_batch_results(batch_id, output_filename="batch_output.jsonl"):
    """
    Download completed batch results.
    """
    batch = client.batches.retrieve(batch_id)
    if batch.status != "completed":
        print(f"Batch not ready yet. Status: {batch.status}")
        return None

    output_file = client.files.retrieve(batch.output_file_id)
    result_content = client.files.content(output_file.id).text

    with open(output_filename, "w") as f:
        f.write(result_content)

    print(f"Results saved to {output_filename}")
    return output_filename

pricing_dict = {
    'gpt-5-mini': {'Input': 0.25, 'Cached input': 0.025, 'Output': 2.00},
    'gpt-5': {'Input': 1.25, 'Cached input': 0.125, 'Output': 10.00},
    'gpt-4.1':  {'Input': 2.00, 'Cached input': 0.5, 'Output': 8.00}
}

def parse_batch_results(output_filename, model, pricing_dict=pricing_dict):
    """
    Parse JSONL batch output file into DataFrame of responses,
    including token usage and cached inputs for cost calculation.
    Uses pricing_dict to compute estimated cost (values are $ per 1M tokens).
    """
    responses = []
    total_prompt_tokens = 0
    total_completion_tokens = 0
    total_cached_tokens = 0

    with open(output_filename, "r") as f:
        for line in f:
            item = json.loads(line)
            custom_id = item["custom_id"]

            body = item["response"]["body"]
            choice = json.loads(body["choices"][0]["message"]["content"])

            # usage info
            usage = body.get("usage", {})
            prompt_tokens = usage.get("prompt_tokens", 0)
            completion_tokens = usage.get("completion_tokens", 0)
            cached_tokens = (
                usage.get("prompt_tokens_details", {})
                .get("cached_tokens", 0)
            )

            total_prompt_tokens += prompt_tokens
            total_completion_tokens += completion_tokens
            total_cached_tokens += cached_tokens

             # Flatten choice dict into the response row
            row = {
                "custom_id": custom_id,
                "prompt_tokens": prompt_tokens,
                "completion_tokens": completion_tokens,
                "cached_tokens": cached_tokens,
                "total_tokens": usage.get("total_tokens", prompt_tokens + completion_tokens),
            }
            # Add each key from choice dict as a separate column
            row.update(choice)
            responses.append(row)

    df = pd.DataFrame(responses)

    # === Token summary ===
    print("=== Token Usage Summary ===")
    print(f"Prompt tokens:     {total_prompt_tokens}")
    print(f"Completion tokens: {total_completion_tokens}")
    print(f"Cached tokens:     {total_cached_tokens}")
    print(f"Total tokens:      {total_prompt_tokens + total_completion_tokens}")

    # === Cost calculation ===
    if model not in pricing_dict:
        raise ValueError(f"Model {model} not found in pricing_dict")

    # Convert $ per 1M tokens to $ per token
    rates = pricing_dict[model]
    prompt_rate = rates["Input"] / 1_000_000
    cached_rate = rates["Cached input"] / 1_000_000
    completion_rate = rates["Output"] / 1_000_000

    cost = (
        (total_prompt_tokens - total_cached_tokens) * prompt_rate +
        (total_cached_tokens * cached_rate) +
        (total_completion_tokens * completion_rate)
    ) / 2 # 50% for batch processing

    print(f"Estimated cost for {model}: ${cost:.4f}")
    return df

batch_data_path = os.path.join("..", "..", "data", "processed", "batch_data")

# --- 1. Prepare and split DataFrame ---
def prepare_batches(df_articles, n_batches, BATCH_INFO_FILE):
    """Split df_articles into 5 non-overlapping batches and create batch_info.pkl.
       If batches already exist, load them instead of recreating."""

    if os.path.exists(BATCH_INFO_FILE):
        print(f"⚠️ {BATCH_INFO_FILE} already exists. Loading existing batch definitions.")
        batch_info = pd.read_pickle(BATCH_INFO_FILE)
        dfs = [pd.read_pickle(os.path.join(batch_data_path, f"batch_{i}_data.pkl")) for i in batch_info["batch_number"]]
        return dfs, batch_info

    # --- Only runs if no prior batch info file found ---
    df_articles = df_articles.reset_index(drop=True)
    #df_articles["row_id"] = range(len(df_articles))
    df_articles = df_articles.sample(frac=1, random_state=42).reset_index(drop=True)

    n = len(df_articles)
    part_size = n // n_batches
    dfs = []

    for i in range(n_batches):
        start_idx = i * part_size
        end_idx = (i + 1) * part_size if i < n_batches - 1 else n
        df_part = df_articles.iloc[start_idx:end_idx].copy()
        df_part["batch_number"] = i + 1
        dfs.append(df_part)

    batch_info = pd.DataFrame({
        "batch_number": range(1, n_batches + 1),
        "n_rows": [len(d) for d in dfs],
        "status": ["pending"] * n_batches,
        "batch_id": [None] * n_batches,
        "input_file": [os.path.join(batch_data_path, f"batch_{i+1}_input.jsonl") for i in range(n_batches)],
        "output_file": [os.path.join(batch_data_path, f"batch_{i+1}_output.jsonl") for i in range(n_batches)],
        "result_file": [os.path.join(batch_data_path, f"batch_{i+1}_results.pkl") for i in range(n_batches)],
    })

    # Save split dataframes
    for i, d in enumerate(dfs, start=1):
        d.to_pickle(os.path.join(batch_data_path, f"batch_{i}_data.pkl"))

    batch_info.to_pickle(BATCH_INFO_FILE)
    print(f"✅ Created {n_batches} batches and saved {BATCH_INFO_FILE}")
    return dfs, batch_info


# --- 2. Run one batch at a time ---
def run_batch(batch_to_run, system_prompt, response_format, model, temperature, reasoning_effort, client, BATCH_INFO_FILE):
    # Load batch info and data
    batch_info = pd.read_pickle(BATCH_INFO_FILE)
    df_sampled = pd.read_pickle(os.path.join(batch_data_path, f"batch_{batch_to_run}_data.pkl"))

    input_file = batch_info.loc[batch_info.batch_number == batch_to_run, "input_file"].iloc[0]
    output_file = batch_info.loc[batch_info.batch_number == batch_to_run, "output_file"].iloc[0]
    result_file = batch_info.loc[batch_info.batch_number == batch_to_run, "result_file"].iloc[0]

    if not os.path.exists(input_file):
        print(f"🚀 Running batch {batch_to_run} with {len(df_sampled)} rows")
        # Step 1: Create batch input
        create_batch(df_sampled, system_prompt, response_format, model, temperature, reasoning_effort, input_file)
        # Step 2: Submit the batch input
        batch_id = submit_batch(batch_input_filename=input_file)
        batch_info.loc[batch_info.batch_number == batch_to_run, "batch_id"] = batch_id
        print(f"📤 Submitted batch {batch_to_run} → Batch ID: {batch_id}")
        batch_status = client.batches.retrieve(batch_id)
        batch_info.loc[batch_info.batch_number == batch_to_run, "status"] = batch_status.status
        batch_info.loc[batch_info.batch_number == batch_to_run, "batch_id"] = batch_id
        batch_info.to_pickle(BATCH_INFO_FILE)

    else:
        batch_id = batch_info.loc[batch_info.batch_number == batch_to_run, "batch_id"].iloc[0]
        print(f"⚠️ Output already exists for batch {batch_to_run}. Skipping submission. Batch ID: {batch_id}")

def check_batch(batch_to_run, BATCH_INFO_FILE, model):
    batch_info = pd.read_pickle(BATCH_INFO_FILE)
    batch_id = batch_info.loc[batch_info.batch_number == batch_to_run, "batch_id"].iloc[0]
    input_file = batch_info.loc[batch_info.batch_number == batch_to_run, "input_file"].iloc[0]
    output_file = batch_info.loc[batch_info.batch_number == batch_to_run, "output_file"].iloc[0]
    result_file = batch_info.loc[batch_info.batch_number == batch_to_run, "result_file"].iloc[0]
    batch_status = batch_info.loc[batch_info.batch_number == batch_to_run, "input_file"].iloc[0]

    # Step 3: Wait/check status
    #print(client.batches.retrieve(batch_id)) # uncomment for more details
    if batch_id:
        new_batch_status = client.batches.retrieve(batch_id).status
        if new_batch_status != batch_status:
            batch_status = new_batch_status
            batch_info.loc[batch_info.batch_number == batch_to_run, "status"] = batch_status
            batch_info.to_pickle(BATCH_INFO_FILE)
            print("Batch status:", batch_status)

        if batch_status == "completed":
            if not os.path.exists(output_file):
                output_file = retrieve_batch_results(batch_id, output_file)
                print("📥 Batch results downloaded.")
            else:
                print("Results arlready downloaded.")

        # Step 4: Parse batch output
        if os.path.exists(output_file):
            df_responses = parse_batch_results(output_file, model)

            # Step 5: Merge with original
            df_sampled = pd.read_pickle(os.path.join(batch_data_path, f"batch_{batch_to_run}_data.pkl")).reset_index(drop = True)
            df_with_results = pd.concat([df_sampled, df_responses], axis=1)
            df_with_results.to_pickle(result_file)
            print(f"✅ Saved merged results to {result_file}")

            batch_info.loc[batch_info.batch_number == batch_to_run, "status"] = "completed"
            batch_info.to_pickle(BATCH_INFO_FILE)
            print("🗂️ Updated batch_info.pkl")

        else:
            print(f"⚠️ Output file not found for batch {batch_to_run}")
    else:
        print("No batch id found.")

# --- 3. Combine all completed results ---
def combine_all_batches(BATCH_INFO_FILE):
    batch_info = pd.read_pickle(BATCH_INFO_FILE)
    completed = batch_info[batch_info.status == "completed"]

    if completed.empty:
        print("⚠️ No completed batches to combine.")
        return None

    dfs = []
    for result_file in completed["result_file"]:
        if os.path.exists(result_file):
            dfs.append(pd.read_pickle(result_file))

    if dfs:
        df_all = pd.concat(dfs, ignore_index=True)
        #df_all.to_pickle("all_batches_combined.pkl")
        print(f"✅ Combined {len(dfs)} completed batches")# → all_batches_combined.pkl")
        return df_all
    else:
        print("⚠️ No result files found.")
        return None

# save examples for folder
def save_examples_to_folder(var, df_articles_with_results, examples_dir, var_original = None):

    var_abbr = var.removesuffix("_stand_flat").removesuffix("_stand")

    base_folder = os.path.join(examples_dir, var_abbr)
    if not os.path.exists(base_folder):
        os.makedirs(base_folder)

    # 2️⃣ Iterate over each unique value in race_ethnicity_stand_flat
    for val in set(df_articles_with_results[var]):
        # Create subfolder for this race/ethnicity if it doesn't exist
        subfolder_path = os.path.join(base_folder, str(val).replace("/","-"))
        if not os.path.exists(subfolder_path):
            os.makedirs(subfolder_path)

        # 3️⃣ Select 3 New York Times and 3 Other publisher articles for this category
        nyt_articles = df_articles_with_results[
            (df_articles_with_results[var] == val) &
            (df_articles_with_results['publisher'] == "New York Times")
        ].head(3)

        other_articles = df_articles_with_results[
            (df_articles_with_results[var] == val) &
            (df_articles_with_results['publisher'] != "New York Times")
        ].head(3)

        # Combine both selections
        selected_articles = pd.concat([nyt_articles, other_articles], ignore_index=True)

        # 4️⃣ Save each article as a .docx file
        for n, row in enumerate(selected_articles.itertuples(), start=1):
            doc = Document()

            # Add fields to document
            doc.add_heading(str(row.title), level=1)
            doc.add_paragraph(f"Publisher: {row.publisher}")
            doc.add_paragraph(f"Date: {row.date}")
            doc.add_paragraph(f"Section: {row.section}")
            doc.add_paragraph(f"Source file: {row.source_file}")
            if var_original is None:
                doc.add_paragraph(f"Raw {var_abbr.replace("_", " ")} result, '{var_abbr}': {getattr(row, var_abbr)}")
                doc.add_paragraph(f"Standardized {var_abbr.replace("_", " ")} result, '{var}': {getattr(row, var)}")
            else:
                doc.add_paragraph(f"Raw {var_original.replace("_", " ")} result, '{var_original}': {getattr(row, var_original)}")
                doc.add_paragraph(f"Standardized {var_original.replace("_", " ")} result, '{var}': {getattr(row, var)}")
            doc.add_paragraph("\nBody:\n")
            doc.add_paragraph(str(row.body))

            # Create filename
            safe_publisher = row.publisher.replace(" ", "_")
            filename = f"{n}.{var_abbr}_{str(val).replace("/","-")}_{safe_publisher}.docx"
            filepath = os.path.join(subfolder_path, filename)

            # Save document
            doc.save(filepath)

    print("✅ Articles saved successfully.")

# classifying other occupations
def create_batch_other(unclassified_items, occupation_categories, batch_input_filename = "classify_other_occupation_batch.jsonl"):
    with open(batch_input_filename, "w") as f:
        for i, sublist in enumerate(unclassified_items):
            if sublist:
                request = {
                    "custom_id": f"classify_{i}",
                    "method": "POST",
                    "url": "/v1/chat/completions",
                    "body": {
                        "model": "gpt-5",
                        "messages": [
                            {"role": "system", "content": (
                                "Classify each user-provided occupation into one of the following categories, using 'Other' sparingly only if item cannot be classified. "
                                f"Occupation categories: {occupation_categories}"
                            )},
                            {"role": "user", "content": f"{sublist}"}
                        ],
                        "reasoning_effort": 'minimal',
                        "response_format": {
                            "type": "json_schema",
                            "json_schema": {
                                "name": "occupation_category_list",
                                "schema": {
                                    "type": "object",
                                    "properties": {
                                        "categories": {
                                            "type": "array",
                                            "items": {"type": "string"},
                                        }
                                    },
                                    "required": ["categories"],
                                },
                            },
                        },
                    }
                }
                f.write(json.dumps(request) + "\n")
    return batch_input_filename

def standardize_occupation(values_list, categories):
    stand_list = []
    unclassified_list = []

    if values_list == ["NA"]:
        return [], []

    for value in values_list:
        cat_found = False

        for cat in categories:
            if cat.lower() in value.lower():
                stand_list.append(cat)
                cat_found = True
                break

        if not cat_found:
            if "administrative" in value.lower() or "clerical" in value.lower():
                stand_list.append('Administrative/Clerical')

            elif "unspecified" not in value.lower() and "no specific" not in value.lower() and "general" not in value.lower() and "not specific" not in value.lower():
                unclassified_list.append(value)
    # Remove duplicates from stand_list
    stand_list = list(set(stand_list))

    return stand_list, unclassified_list


# delete??
def classify_other(other_list, occupation_categories):

    response_format = {"type": "json_schema",
                      "json_schema":
                        {
                      "name": "occupation_category_list",
                      "schema": {
                        "type": "object",
                        "properties": {
                          "categories": {
                            "type": "array",
                            "description": "List of occupation categories, one per user-provided occupation, corresponding in order. Only allowed values: 'Service', 'Manufacturing', 'Construction', 'Administrative/Clerical', 'Transportation', 'Farming', or 'Other'.",
                            "items": {
                              "type": "string",
                              "enum": [
                                "Service",
                                "Manufacturing",
                                "Construction",
                                "Administrative/Clerical",
                                "Transportation",
                                "Farming",
                                "Other"
                              ]
                            }
                          }
                        },
                        "required": [
                          "categories"
                        ],
                        "additionalProperties": False
                      },
                      "strict": True
                    }}


    response = client.chat.completions.create(
        model= "gpt-5",
        messages=[
            {"role": "system", "content":
                    ("Classify each user-provided occupation into one of the following categories, using 'Other' sparingly only if item cannot be classified. Output only the category without any rationale, reasoning, or justification. "
                     f"Occupation categories: {occupation_categories}")
            },
            {"role": "user", "content": f"{other_list}"}
        ],
        temperature=temperature,
        reasoning_effort=reasoning_effort,
        response_format = response_format
    )

    return response.choices[0].message.content
