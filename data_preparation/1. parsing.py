from docx import Document
import pandas as pd
from datetime import datetime
import os
import time
import glob
import argparse
from pathlib import Path

def clean_date(cleaned_date_str):
    """
    Convert a string date in the format 'Month Day, Year' (e.g., 'November 19, 2024')
    into the ISO format 'YYYY-MM-DD'.
    """
    parsed_date = datetime.strptime(cleaned_date_str, '%B %d, %Y')
    return parsed_date.strftime('%Y-%m-%d')

def extract_articles_to_dataframe(doc_path, save_intermediate=False):
    """
    Parse a .docx file and extract structured article data into a pandas DataFrame.
    Articles are identified by 'Heading 1' style and include metadata like title, publisher,
    date, section, length, and body. The function also handles date corrections and load dates.
    Optionally saves intermediate DataFrame as a pickle file.
    """
    # Start timing the process
    start_time = time.time()
    print("Processing document: ", doc_path)

    # Load the .docx file
    doc = Document(doc_path)
    count = 0
    articles = []

    # Convert path to pathlib object and generate relative file path
    doc_path = Path(doc_path)
    relative_path = f"{doc_path.parent.name}/{doc_path.name}"

    # Count total paragraphs for iteration
    total_paragraphs = len(doc.paragraphs)

    # Iterate through the paragraphs
    while count < total_paragraphs:
        para = doc.paragraphs[count]
        count += 1

        # Detect the beginning of an article using the 'Heading 1' style
        if "Heading 1" in para.style.name:
            article = {}
            correction_appended = False

            # Extract basic metadata
            article['title'] = para.text.strip()
            article['publisher'] =  doc.paragraphs[count].text
            raw_date = doc.paragraphs[count+1].text

            # Remove trailing weekday if present (e.g., "Tuesday")
            date_parts = raw_date.split()[:3]  # "November 19, 2024"
            cleaned_date_str = " ".join(date_parts).strip(",")
            load_date_at_end = False

            # Try parsing the cleaned date
            try:
                article['date'] = clean_date(cleaned_date_str)
            except:
                # If date parsing fails, check for correction notice or defer to later
                if doc.paragraphs[count+2].text.strip("\n ") == "Correction Appended":
                    correction_appended = True
                else:
                    raw_date = doc.paragraphs[count+1].text
                    date_parts = raw_date.split()[:3]
                    try:
                        article['date'] = clean_date(cleaned_date_str)
                    except:
                        load_date_at_end = True

            # Look ahead for 'Section:', 'Length:', and 'Body' markers
            for i in range(3,20):
                text = doc.paragraphs[count+i].text
                if text[:8] == "Section:":
                    section = text[9:]
                    article["section"] = section
                if text[:7] == "Length:":
                    article["length"] = text[8:]
                if text == "Body":
                    count_body_start = count + i + 1
                    break

            # Extract article body until "End of Document"
            body = []
            new_count = count_body_start
            while text.strip("\n ").lower() != "end of document":
                text = doc.paragraphs[new_count].text
                body.append(text)
                new_count += 1
            article["body"] = "\n".join(body[:-1])

            # If correction was appended, extract the corrected date from body
            if correction_appended:
                date_parts  = article["body"].split("Correction-Date: ")[-1].split()[:3]
                cleaned_date_str = " ".join(date_parts).strip(",")
                article['date'] = clean_date(cleaned_date_str)

            # If date wasn't parsed earlier, attempt to extract Load-Date from body
            if load_date_at_end:
                try:
                    date_parts  = article["body"].split("Load-Date: ")[-1].split()[:3]
                    cleaned_date_str = " ".join(date_parts).strip(",")
                    article['date'] = clean_date(cleaned_date_str)
                except:
                    print(f"Date not found in doc {relative_path}. See: {article['title']}.")

            # Update counter to move to next article
            count = new_count

            # Add correction and fallback flags
            article["correction_appended"] = correction_appended
            article["load_date_at_end"] = load_date_at_end

            # Store the article
            articles.append(article)

    print("Number of articles: ", len(articles))

    # Create a DataFrame from all parsed articles
    df = pd.DataFrame(articles)
    df["source_file"] = relative_path

    # Optionally save intermediate results
    if save_intermediate:
        # Ensure processed directory exists
        output_folder = Path("../data/processed/parsed_files/") # temp
        output_folder.mkdir(parents=True, exist_ok=True) # temp

        df.to_pickle(output_folder / f"{os.path.splitext(os.path.basename(doc_path))[0]}_parsed_articles.pkl")

     # End timing and report duration
    end_time = time.time()
    elapsed_time = end_time - start_time
    print(f"Elapsed time: {elapsed_time / 60:.2f} minutes")

    return df

# Entry point for command-line usage
if __name__ == "__main__":
    from functools import partial
    from concurrent.futures import ProcessPoolExecutor

    # Define command-line arguments
    parser = argparse.ArgumentParser(description="Parse DOCX files into DataFrames.")
    parser.add_argument('--save-intermediate', action='store_true',
                        help='Save individual parsed articles as pickles.')
    args = parser.parse_args()

    # Find all DOCX files in the raw data folder
    docx_files = glob.glob(os.path.join('..', 'data', 'raw', '*', '*.DOCX'))
    print("Files to process: ", len(docx_files))

    # Partially apply the save_intermediate argument for use in parallel
    process_func = partial(extract_articles_to_dataframe, save_intermediate=args.save_intermediate)

    # Use multiprocessing to process all files in parallel
    with ProcessPoolExecutor(max_workers=os.cpu_count() - 1) as executor:
        results = list(executor.map(process_func, docx_files))

    # Concatenate all resulting DataFrames
    df_articles = pd.concat(results, ignore_index=True)

    # Ensure output directory exists
    output_folder = Path("../data/processed/")
    output_folder.mkdir(parents=True, exist_ok=True)

    # Save the final combined DataFrame to a pickle file
    df_articles.to_pickle(output_folder / "parsed_articles.pkl")

    print("Done.")
